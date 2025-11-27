from fastapi import FastAPI
from pydantic import BaseModel
from typing import List
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import uuid
from fastapi.staticfiles import StaticFiles

# ==================== CONFIGURAÇÃO ====================

app = FastAPI(title="API BIQ ADV Farma")

OUTPUT_DIR = "generated"
os.makedirs(OUTPUT_DIR, exist_ok=True)

app.mount("/files", StaticFiles(directory=OUTPUT_DIR), name="files")

MODELO_PATH = "MODELO_BIQ.docx"

# ==================== MODELOS DE DADOS ====================

class Acao(BaseModel):
    numero: str
    tipo: str
    acao: str
    explicacao: str
    responsavel: str
    local: str
    data: str

class Membro(BaseModel):
    nome: str
    cargo: str

class BIQData(BaseModel):
    dados_incidentes: dict
    justificativa_texto: str
    descricao_incidente: str
    acoes: List[List[str]]
    equipe: List[List[str]]
    follow_up_texto: str
    numero_biq: str

# ==================== ROTA PRINCIPAL ====================

@app.post("/gerar-biq-docx")
def gerar_biq_docx(biq: BIQData):
    doc = Document(MODELO_PATH)

    dados_incidentes = biq.dados_incidentes
    justificativa_texto = biq.justificativa_texto
    descricao_incidente = biq.descricao_incidente
    acoes = biq.acoes
    equipe = biq.equipe
    follow_up_texto = biq.follow_up_texto

    # ======== PREENCHIMENTO DAS TABELAS ========

    for table in doc.tables:
        header = table.rows[0].cells[0].text.strip()

        # Dados da incidência
        if "Campo" in header:
            for row in table.rows[1:]:
                campo = row.cells[0].text.strip()
                if campo in dados_incidentes:
                    row.cells[1].text = dados_incidentes[campo]

        # Justificativa da classificação
        elif "Justificativa" in header:
            for row in table.rows[1:]:
                if "Justificativa" in row.cells[0].text:
                    row.cells[1].text = justificativa_texto
                elif "Reincidência" in row.cells[0].text:
                    row.cells[1].text = "Não."

        # Descrição da incidência
        elif "Descrição" in header:
            for row in table.rows:
                if "Descrição" in row.cells[0].text:
                    row.cells[1].text = descricao_incidente

        # Ações corretivas e preventivas
        elif "N° Ação" in header:
            while len(table.rows) > 1:
                table._element.remove(table.rows[1]._element)
            for acao in acoes:
                row = table.add_row().cells
                for i, valor in enumerate(acao):
                    row[i].text = valor

        # Tabela de anexos
        elif "Número do Anexo" in header:
            while len(table.rows) > 1:
                table._element.remove(table.rows[1]._element)
            row = table.add_row().cells
            row[0].text = "Não Aplicável"
            row[1].text = "Não Aplicável"
            row[2].text = "Não Aplicável"

        # Envolvimento da equipe
        elif "Nome" in header and "Cargo" in table.rows[0].cells[1].text:
            while len(table.rows) > 1:
                table._element.remove(table.rows[1]._element)
            for membro in equipe:
                row = table.add_row().cells
                row[0].text = membro[0]
                row[1].text = membro[1]
                row[2].text = "___________________________"

        # Follow-up
        elif "Follow" in header:
            for row in table.rows:
                if "Follow" in row.cells[0].text:
                    row.cells[1].text = follow_up_texto

    # ======== CABEÇALHO ========

    for section in doc.sections:
        header = section.header
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "BIQ" in cell.text.strip():
                        cell.text = biq.numero_biq
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.name = "Arial"
                                run.font.size = Pt(10)

    # ======== PADRONIZAÇÃO DE FONTE ========

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(9)

    # ======== SALVAR DOCX ========

    filename = f"BIQ_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)

    return {"file_url": f"/files/{filename}"}

# ==================== TESTE ====================
@app.get("/")
def home():
    return {"status": "ok", "msg": "API de preenchimento automático do BIQ ADV Farma"}
