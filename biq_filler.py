# -*- coding: utf-8 -*-
"""
Módulo responsável por preencher o Boletim de Incidência da Qualidade (Anexo 01 POP-NO-GQ-157)
em conformidade com as BPF e formatação padrão ADV Farma.
"""

from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ======= PADRÕES DE FORMATAÇÃO =======
FONT_NAME = "Arial"
FONT_SIZE_PT = 9


def _format_cell(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = FONT_NAME
            r.font.size = Pt(FONT_SIZE_PT)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def preencher_biq_from_payload(template_path: str, payload: dict) -> bytes:
    """Preenche o modelo oficial do BIQ e retorna o binário do DOCX."""
    doc = Document(template_path)
    dados = payload.get("dados_incidentes", {})
    justificativa = payload.get("justificativa_texto", "")
    descricao = payload.get("descricao_incidente", "")
    acoes = payload.get("acoes", [])
    equipe = payload.get("equipe", [])
    follow = payload.get("follow_up_texto", "")
    numero_biq = payload.get("numero_biq", "—")

    # ========== PREENCHIMENTO DAS TABELAS ==========
    for table in doc.tables:
        header = table.rows[0].cells[0].text.strip()

        if "Campo" in header:
            for row in table.rows[1:]:
                campo = row.cells[0].text.strip()
                if campo in dados:
                    row.cells[1].text = dados[campo]
                    _format_cell(row.cells[1])

        elif "Justificativa" in header:
            for row in table.rows[1:]:
                if "Justificativa" in row.cells[0].text:
                    row.cells[1].text = justificativa
                elif "Reincidência" in row.cells[0].text:
                    row.cells[1].text = "Não."
                _format_cell(row.cells[1])

        elif "Descrição" in header:
            for row in table.rows:
                if "Descrição" in row.cells[0].text:
                    row.cells[1].text = descricao
                    _format_cell(row.cells[1])

        elif "N° Ação" in header:
            while len(table.rows) > 1:
                table._element.remove(table.rows[1]._element)
            for acao in acoes:
                row = table.add_row().cells
                for i, valor in enumerate(acao):
                    row[i].text = valor
                    _format_cell(row[i])

        elif "Número do Anexo" in header:
            while len(table.rows) > 1:
                table._element.remove(table.rows[1]._element)
            row = table.add_row().cells
            row[0].text = "Não Aplicável"
            row[1].text = "Não Aplicável"
            row[2].text = "Não Aplicável"
            _format_cell(row[0])

        elif "Nome" in header and "Cargo" in table.rows[0].cells[1].text:
            while len(table.rows) > 1:
                table._element.remove(table.rows[1]._element)
            for membro in equipe:
                row = table.add_row().cells
                row[0].text = membro[0]
                row[1].text = membro[1]
                row[2].text = "___________________________"
                _format_cell(row[0])

        elif "Follow" in header:
            for row in table.rows:
                if "Follow" in row.cells[0].text:
                    row.cells[1].text = follow
                    _format_cell(row.cells[1])

    # ========== CABEÇALHO ==========
    for section in doc.sections:
        header = section.header
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "BIQ" in cell.text.strip():
                        cell.text = numero_biq
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for r in p.runs:
                                r.bold = True
                                r.font.name = FONT_NAME
                                r.font.size = Pt(10)

    # ========== SALVA EM MEMÓRIA ==========
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
